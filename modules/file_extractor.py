"""
Module 2: File Text Extraction Engine
Extracts text from PDF and DOCX files for TTS processing.

Supported formats:
- PDF (single/multi-page)
- DOCX (Word documents)
"""

from typing import Optional, List
from dataclasses import dataclass
from pathlib import Path
import io

# PDF extraction libraries
import PyPDF2
import pdfplumber

# Word document library
from docx import Document

from loguru import logger


@dataclass
class ExtractionResult:
    """Result of file text extraction"""
    success: bool
    text: str
    file_path: str
    file_type: str
    page_count: int
    char_count: int
    error: Optional[str] = None


class FileExtractor:
    """Extract text from PDF and DOCX files"""
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}
    MAX_FILE_SIZE_MB = 10
    
    def __init__(self):
        logger.info("FileExtractor initialized")
    
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Main extraction method - routes to specific extractor
        
        Args:
            file_path: Path to file (PDF or DOCX)
            
        Returns:
            ExtractionResult with extracted text
        """
        try:
            path = Path(file_path)
            
            # Validate file
            error = self._validate_file(path)
            if error:
                return self._create_error_result(str(path), error)
            
            # Route to appropriate extractor
            extension = path.suffix.lower()
            
            if extension == '.pdf':
                return self._extract_pdf(path)
            elif extension == '.docx':
                return self._extract_docx(path)
            else:
                return self._create_error_result(
                    str(path), 
                    f"Unsupported file type: {extension}"
                )
        
        except Exception as e:
            logger.error(f"Extraction error: {str(e)}")
            return self._create_error_result(file_path, f"Extraction failed: {str(e)}")
    
    
    def _validate_file(self, path: Path) -> Optional[str]:
        """Validate file exists and is supported"""
        
        if not path.exists():
            return f"File not found: {path}"
        
        if not path.is_file():
            return f"Not a file: {path}"
        
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            return f"Unsupported file type: {extension}. Supported: {self.SUPPORTED_EXTENSIONS}"
        
        # Check file size
        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > self.MAX_FILE_SIZE_MB:
            return f"File too large: {size_mb:.2f}MB (max: {self.MAX_FILE_SIZE_MB}MB)"
        
        return None
    
    
    def _extract_pdf(self, path: Path) -> ExtractionResult:
        """
        Extract text from PDF using pdfplumber (primary) with PyPDF2 fallback
        
        Args:
            path: Path to PDF file
            
        Returns:
            ExtractionResult
        """
        logger.info(f"Extracting PDF: {path.name}")
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            text, page_count = self._extract_pdf_pdfplumber(path)
            
            if not text.strip():
                # Fallback to PyPDF2
                logger.warning("pdfplumber returned empty text, trying PyPDF2")
                text, page_count = self._extract_pdf_pypdf2(path)
            
            if not text.strip():
                return self._create_error_result(
                    str(path), 
                    "No text found in PDF (might be scanned image)"
                )
            
            return ExtractionResult(
                success=True,
                text=text.strip(),
                file_path=str(path),
                file_type='pdf',
                page_count=page_count,
                char_count=len(text.strip()),
                error=None
            )
        
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return self._create_error_result(str(path), f"PDF extraction error: {str(e)}")
    
    
    def _extract_pdf_pdfplumber(self, path: Path) -> tuple[str, int]:
        """Extract using pdfplumber (better for tables and complex layouts)"""
        
        text_parts = []
        page_count = 0
        
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n\n'.join(text_parts), page_count
    
    
    def _extract_pdf_pypdf2(self, path: Path) -> tuple[str, int]:
        """Extract using PyPDF2 (fallback method)"""
        
        text_parts = []
        
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            page_count = len(reader.pages)
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n\n'.join(text_parts), page_count
    
    
    def _extract_docx(self, path: Path) -> ExtractionResult:
        """
        Extract text from Word document
        
        Args:
            path: Path to DOCX file
            
        Returns:
            ExtractionResult
        """
        logger.info(f"Extracting DOCX: {path.name}")
        
        try:
            doc = Document(path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            if not paragraphs:
                return self._create_error_result(
                    str(path), 
                    "No text found in document"
                )
            
            text = '\n\n'.join(paragraphs)
            
            return ExtractionResult(
                success=True,
                text=text.strip(),
                file_path=str(path),
                file_type='docx',
                page_count=1,  # DOCX doesn't have clear page boundaries
                char_count=len(text.strip()),
                error=None
            )
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return self._create_error_result(str(path), f"DOCX extraction error: {str(e)}")
    
    
    def extract_from_bytes(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """
        Extract text from file bytes (for file uploads in web apps)
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (to determine type)
            
        Returns:
            ExtractionResult
        """
        extension = Path(filename).suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_pdf_from_bytes(file_bytes, filename)
            elif extension == '.docx':
                return self._extract_docx_from_bytes(file_bytes, filename)
            else:
                return self._create_error_result(filename, f"Unsupported file type: {extension}")
        
        except Exception as e:
            return self._create_error_result(filename, f"Extraction error: {str(e)}")
    
    
    def _extract_pdf_from_bytes(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """Extract PDF from bytes"""
        
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        text = '\n\n'.join(text_parts)
        
        if not text.strip():
            return self._create_error_result(filename, "No text found in PDF")
        
        return ExtractionResult(
            success=True,
            text=text.strip(),
            file_path=filename,
            file_type='pdf',
            page_count=page_count,
            char_count=len(text.strip()),
            error=None
        )
    
    
    def _extract_docx_from_bytes(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """Extract DOCX from bytes"""
        
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        if not paragraphs:
            return self._create_error_result(filename, "No text found in document")
        
        text = '\n\n'.join(paragraphs)
        
        return ExtractionResult(
            success=True,
            text=text.strip(),
            file_path=filename,
            file_type='docx',
            page_count=1,
            char_count=len(text.strip()),
            error=None
        )
    
    
    def _create_error_result(self, file_path: str, error: str) -> ExtractionResult:
        """Create error result"""
        logger.warning(f"Extraction failed for {file_path}: {error}")
        
        return ExtractionResult(
            success=False,
            text="",
            file_path=file_path,
            file_type="unknown",
            page_count=0,
            char_count=0,
            error=error
        )


# Convenience function
def extract_text(file_path: str) -> ExtractionResult:
    """
    Quick extraction function
    
    Example:
        >>> from modules.file_extractor import extract_text
        >>> result = extract_text("document.pdf")
        >>> print(result.text)
    """
    extractor = FileExtractor()
    return extractor.extract(file_path)


if __name__ == "__main__":
    print("="*60)
    print("MODULE 2: File Extractor - Self Test")
    print("="*60)
    print("\n⚠️  No test files found. Create sample files to test:")
    print("   1. Create 'sample_files' folder")
    print("   2. Add test PDF/DOCX files")
    print("   3. Run: python demo_file_extractor.py")